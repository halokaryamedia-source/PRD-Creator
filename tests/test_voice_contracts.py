from __future__ import annotations

import html
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BUILDER = ROOT / "kits" / "voice-production-kit" / "builder" / "build_docx.py"
VALIDATOR = ROOT / "kits" / "voice-production-kit" / "validator" / "validate.py"


def run_cli(*args: Path | str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, *(str(arg) for arg in args)],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )


def requirements(extra_id: bool = False, type_override: str | None = None) -> str:
    intro_type = type_override or "Main Story"
    text = f"""# Contract Fixture Voice Requirements

## Intro

### VO-INTRO-01 — Welcome

- Type: {intro_type}
- Function: briefing
- Necessity: required
- Speaker: Narrator
- Channel: Direct
- Trigger: Trial start before active play begins.
- Purpose: Tell the player to begin the trial.
- Timing Constraint: Must fit the 3-second opening slot.
- Must communicate:
  - Begin the trial.
- Must not add/repeat:
  - No additional project facts.
- Source refs:
  - Contract fixture.

## Ending

### VO-END-01 — Complete

- Type: Direct NPC Dialogue
- Function: completion
- Necessity: required
- Speaker: Guide
- Channel: Direct
- Trigger: Trial completion after the final objective resolves.
- Purpose: Acknowledge that the trial is complete.
- Must communicate:
  - The trial is complete.
- Must not add/repeat:
  - No additional reward.
- Source refs:
  - Contract fixture.
"""
    if extra_id:
        text += """
## Extra

### VO-EXTRA-01 — Unsupported Extra

- Type: Main Story
- Function: reminder
- Necessity: supporting
- Speaker: Narrator
- Channel: Direct
- Trigger: Never
- Purpose: Exercise the missing-ID parity failure.
- Must communicate:
  - This requirement is intentionally missing from the script.
- Must not add/repeat:
  - Nothing.
- Source refs:
  - Contract fixture.
"""
    return text


SCRIPT = """# Contract Fixture Voice Production
Version: 1.0
Source Voice Requirements: work/voice-requirements.md

Voice Cast:
- Narrator: William Shanks - Rich and Deep
- Guide: Clara - Calm and Clear

## Intro

### VO-INTRO-01 — Welcome
Type: Main Story
Speaker: Narrator
Estimated Duration: 2–3 seconds
```performance
[calm]
Begin the trial.
```

## Ending

### VO-END-01 — Complete
Type: Direct NPC Dialogue
Speaker: Guide
Estimated Duration: 2–3 seconds
```performance
[clear]
The trial is complete.
```
"""


def project_html(*, omit_intro_context: bool = False) -> str:
    intro_trigger = "" if omit_intro_context else html.escape("Trial start before active play begins.")
    end_trigger = html.escape("Trial completion after the final objective resolves.")
    return f"""<!doctype html>
<html><head><style id="production-assets-style"></style></head><body>
<div class="nav-group production-assets-nav">Production Assets</div>
<section data-page-role="production-assets"><div class="voice-objective-shell"></div>
<div class="voice-script-position">Introduction · Voice Line 1/1</div>
<p class="voice-script-context">{intro_trigger}</p>
<pre class="voice-script-text" id="voice-prompt-vo-intro-01">[calm]\nBegin the trial.</pre></section>
<section data-page-role="production-assets"><div class="voice-objective-shell"></div>
<div class="voice-script-position">Ending · Voice Line 1/1</div>
<p class="voice-script-context">{end_trigger}</p>
<pre class="voice-script-text" id="voice-prompt-vo-end-01">[clear]\nThe trial is complete.</pre></section>
</body></html>"""


class VoiceProductionContracts(unittest.TestCase):
    def make_project(self, requirements_text: str | None = None, script_text: str = SCRIPT) -> Path:
        temp = tempfile.TemporaryDirectory()
        self.addCleanup(temp.cleanup)
        project = Path(temp.name)
        (project / "work").mkdir(parents=True)
        (project / "output").mkdir(parents=True)
        (project / "state").mkdir(parents=True)
        (project / "work" / "voice-requirements.md").write_text(
            requirements_text if requirements_text is not None else requirements(),
            encoding="utf-8",
        )
        (project / "work" / "voice-production.md").write_text(script_text, encoding="utf-8")
        (project / "state" / "voice-state.yaml").write_text(
            "status: voice_script_ready\nrevision: contract-1\n",
            encoding="utf-8",
        )
        return project

    def build(self, project: Path) -> subprocess.CompletedProcess[str]:
        return run_cli(
            BUILDER,
            project / "work/voice-production.md",
            project / "output/Voice Production.docx",
            "--requirements",
            project / "work/voice-requirements.md",
        )

    def test_builder_and_validator_happy_path_preserves_section_break_contract(self) -> None:
        project = self.make_project()

        built = self.build(project)
        self.assertEqual(built.returncode, 0, built.stderr or built.stdout)

        validated = run_cli(VALIDATOR, project)
        self.assertEqual(validated.returncode, 0, validated.stderr or validated.stdout)
        self.assertIn("VOICE VALIDATION PASS", validated.stdout)

        doc = Document(project / "output/Voice Production.docx")
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("Speaker: Narrator", full_text)
        self.assertIn("Speaker: Guide", full_text)

        headings = [
            paragraph
            for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 1"
        ]
        self.assertEqual([paragraph.text for paragraph in headings], ["Intro", "Ending"])
        self.assertIsNot(headings[0].paragraph_format.page_break_before, True)
        self.assertIs(headings[1].paragraph_format.page_break_before, True)

    def test_validator_accepts_current_project_html_objective_contract(self) -> None:
        project = self.make_project()
        (project / "output/final.html").write_text(project_html(), encoding="utf-8")

        validated = run_cli(VALIDATOR, project)
        self.assertEqual(validated.returncode, 0, validated.stderr or validated.stdout)
        self.assertIn("project_html=passed", validated.stdout)

    def test_validator_rejects_missing_flow5_trigger_context_in_project_html(self) -> None:
        project = self.make_project()
        (project / "output/final.html").write_text(
            project_html(omit_intro_context=True),
            encoding="utf-8",
        )

        validated = run_cli(VALIDATOR, project)
        self.assertEqual(validated.returncode, 1)
        self.assertIn(
            "Project HTML missing Flow 5 Trigger context for VO-INTRO-01",
            validated.stdout,
        )

    def test_builder_rejects_missing_voice_id_parity(self) -> None:
        project = self.make_project(requirements(extra_id=True))

        built = self.build(project)
        self.assertEqual(built.returncode, 2)
        self.assertIn("Voice requirement parity failed", built.stderr)
        self.assertIn("missing script IDs: VO-EXTRA-01", built.stderr)

    def test_builder_rejects_type_mismatch(self) -> None:
        project = self.make_project(requirements(type_override="Direct NPC Dialogue"))

        built = self.build(project)
        self.assertEqual(built.returncode, 2)
        self.assertIn(
            "Voice Type differs from Flow 5 requirement for: VO-INTRO-01",
            built.stderr,
        )

    def test_builder_rejects_speaker_mismatch(self) -> None:
        script = SCRIPT.replace("Speaker: Narrator", "Speaker: Guide", 1)
        project = self.make_project(script_text=script)

        built = self.build(project)
        self.assertEqual(built.returncode, 2)
        self.assertIn(
            "Voice Speaker differs from Flow 5 requirement for: VO-INTRO-01",
            built.stderr,
        )

    def test_builder_rejects_empty_section_without_traceback(self) -> None:
        script = SCRIPT.replace("## Ending", "## Empty Section\n\n## Ending", 1)
        project = self.make_project(script_text=script)

        built = self.build(project)
        self.assertEqual(built.returncode, 2)
        self.assertNotIn("Traceback", built.stderr)
        self.assertIn("Voice section has no entries: Empty Section", built.stderr)
        self.assertFalse((project / "output/Voice Production.docx").exists())

    def test_validator_rejects_voice_without_initial_performance_tag(self) -> None:
        script = SCRIPT.replace("[calm]\nBegin the trial.", "Begin the trial.", 1)
        project = self.make_project(script_text=script)

        validated = run_cli(VALIDATOR, project)
        self.assertEqual(validated.returncode, 2)
        self.assertIn(
            "VO-INTRO-01 performance must begin with at least one initial [performance direction] tag",
            validated.stderr,
        )


if __name__ == "__main__":
    unittest.main()
