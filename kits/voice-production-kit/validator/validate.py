#!/usr/bin/env python3
from __future__ import annotations
import argparse, hashlib, html, json, re, sys
from dataclasses import dataclass, field
from pathlib import Path
from docx import Document

PLACEHOLDER_RE = re.compile(r"\b(?:TBD|TODO|FIXME)\b|\[OPEN\]", re.I)
ENTRY_RE = re.compile(r"^###\s+([A-Za-z0-9][A-Za-z0-9-]*)\s+[—-]\s+(.+?)\s*$")
VOICE_ID_RE = re.compile(r"\bVO-[A-Z0-9][A-Z0-9-]*\b")
STATUS_RE = re.compile(r"^\s*status:\s*([A-Za-z0-9_-]+)\s*$", re.M)
PERFORMANCE_TAG_LINE_RE = re.compile(r"^(?:\[[^\[\]\r\n]+\]\s*)+$")
SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
SOURCE_PRD_REVISION_RE = re.compile(r"(?mi)^\s*Source PRD revision:\s*(\S+)\s*$")
SOURCE_REQUIREMENTS_RE = re.compile(
    r"(?mi)^\s*Source Voice Requirements:\s*(\S+)\s*/\s*(.+?)\s*\|\s*sha256:([0-9a-f]{64})\s*$"
)

@dataclass
class Requirement:
    voice_id: str
    title: str
    voice_type: str = ""
    speaker: str = ""
    channel: str = ""
    trigger: str = ""
    must_communicate: list[str] = field(default_factory=list)

@dataclass
class ScriptEntry:
    voice_id: str
    title: str
    voice_type: str = ""
    speaker: str = ""
    duration: str = ""
    performance: str = ""
    section: str = ""


def norm(s: str) -> str:
    return " ".join(s.split())


def scalar_values(text: str, key: str) -> list[str]:
    pattern = re.compile(rf"(?m)^\s*{re.escape(key)}:\s*(.*?)\s*(?:#.*)?$")
    values=[]
    for raw in pattern.findall(text):
        value=raw.strip()
        if len(value)>=2 and value[0]==value[-1] and value[0] in {"'", '"'}:
            value=value[1:-1].strip()
        values.append(value)
    return values


def one_scalar(text: str, key: str, owner: str) -> str:
    values=scalar_values(text,key)
    if len(values)!=1 or not values[0]:
        raise ValueError(f"{owner} must define exactly one non-empty {key}")
    return values[0]


def validate_revision_identity(project: Path, req: Path, scr: Path, state: Path) -> list[str]:
    issues: list[str] = []
    state_text=state.read_text(encoding="utf-8")
    source_handoff=one_scalar(state_text,"source_handoff","voice-state.yaml")
    source_prd_revision=one_scalar(state_text,"source_prd_revision","voice-state.yaml")
    project_html=one_scalar(state_text,"project_html","voice-state.yaml")

    handoff_path=project/source_handoff
    if not handoff_path.is_file():
        return [f"Voice source_handoff does not exist: {source_handoff}"]
    handoff_text=handoff_path.read_text(encoding="utf-8")
    handoff_status=one_scalar(handoff_text,"status","handoff-state.yaml")
    accepted_revision=one_scalar(handoff_text,"accepted_prd_version","handoff-state.yaml")
    if handoff_status!="handoff_ready":
        issues.append(f"Upstream PRD handoff status is {handoff_status!r}, expected 'handoff_ready'")

    render_data_path=project/"work/render-data.json"
    if not render_data_path.is_file():
        issues.append("Current render-data.json is missing for Voice revision identity")
        current_revision=""
    else:
        try:
            data=json.loads(render_data_path.read_text(encoding="utf-8"))
            document=data.get("document") if isinstance(data,dict) else None
            current_revision=str(document.get("version") or "").strip() if isinstance(document,dict) else ""
        except json.JSONDecodeError as exc:
            issues.append(f"Current render-data.json is invalid: {exc}")
            current_revision=""

    req_text=req.read_text(encoding="utf-8")
    req_revisions=SOURCE_PRD_REVISION_RE.findall(req_text)
    if len(req_revisions)!=1:
        issues.append("voice-requirements.md must define exactly one Source PRD revision")
        requirements_revision=""
    else:
        requirements_revision=req_revisions[0].strip()

    script_text=scr.read_text(encoding="utf-8")
    source_matches=SOURCE_REQUIREMENTS_RE.findall(script_text)
    if len(source_matches)!=1:
        issues.append(
            "voice-production.md must define exactly one Source Voice Requirements binding with revision, canonical path, and sha256"
        )
        script_revision=script_path=script_sha=""
    else:
        script_revision,script_path,script_sha=(part.strip() for part in source_matches[0])

    expected_revision=accepted_revision
    identities={
        "voice-state source_prd_revision": source_prd_revision,
        "voice-requirements Source PRD revision": requirements_revision,
        "voice-production Source Voice Requirements revision": script_revision,
        "current render-data document.version": current_revision,
    }
    for label,value in identities.items():
        if value!=expected_revision:
            issues.append(f"{label}={value!r}, expected current accepted PRD revision {expected_revision!r}")

    if script_path and script_path!="work/voice-requirements.md":
        issues.append(
            f"voice-production Source Voice Requirements path={script_path!r}, expected 'work/voice-requirements.md'"
        )
    actual_req_sha=hashlib.sha256(req.read_bytes()).hexdigest()
    if script_sha and (SHA256_RE.fullmatch(script_sha) is None or script_sha!=actual_req_sha):
        issues.append(
            "voice-production Source Voice Requirements sha256 does not match current work/voice-requirements.md bytes"
        )

    expected_html=f"output/v{expected_revision}/prd.html"
    if project_html!=expected_html:
        issues.append(f"voice-state project_html={project_html!r}, expected {expected_html!r}")
    return issues


def has_initial_performance_tag(performance: str) -> bool:
    first = next((line.strip() for line in performance.splitlines() if line.strip()), "")
    return bool(first and PERFORMANCE_TAG_LINE_RE.fullmatch(first))


def parse_requirements(path: Path) -> dict[str, Requirement]:
    text = path.read_text(encoding="utf-8")
    if PLACEHOLDER_RE.search(text):
        raise ValueError("Voice requirements contain unresolved placeholders")
    lines = text.splitlines()
    out: dict[str, Requirement] = {}
    current: Requirement | None = None
    list_mode: str | None = None
    for raw in lines:
        line = raw.rstrip()
        m = ENTRY_RE.match(line)
        if m:
            vid = m.group(1)
            if vid in out:
                raise ValueError(f"Duplicate Voice ID in requirements: {vid}")
            current = Requirement(vid, m.group(2).strip())
            out[vid] = current
            list_mode = None
            continue
        if current is None:
            continue
        if line.startswith("- Type:"):
            current.voice_type = line.split(":",1)[1].strip(); list_mode=None
        elif line.startswith("- Speaker:"):
            current.speaker = line.split(":",1)[1].strip(); list_mode=None
        elif line.startswith("- Channel:"):
            current.channel = line.split(":",1)[1].strip(); list_mode=None
        elif line.startswith("- Trigger:"):
            current.trigger = line.split(":",1)[1].strip(); list_mode=None
        elif line.startswith("- Must communicate:"):
            list_mode = "must"
        elif line.startswith("- Must not add/repeat:") or line.startswith("- Source refs:"):
            list_mode = None
        elif list_mode == "must" and re.match(r"^\s{2,}-\s+", line):
            current.must_communicate.append(re.sub(r"^\s{2,}-\s+", "", line).strip())
    if not out:
        raise ValueError("No Voice IDs found in requirements")
    for r in out.values():
        missing = [k for k,v in {"Type":r.voice_type,"Speaker":r.speaker,"Channel":r.channel,"Trigger":r.trigger}.items() if not v]
        if missing:
            raise ValueError(f"{r.voice_id} missing requirement metadata: {', '.join(missing)}")
    return out


def parse_script(path: Path) -> tuple[list[str], dict[str, ScriptEntry]]:
    text = path.read_text(encoding="utf-8")
    if PLACEHOLDER_RE.search(text):
        raise ValueError("Voice script contains unresolved placeholders")
    lines = text.splitlines()
    sections: list[str] = []
    current_section = ""
    out: dict[str, ScriptEntry] = {}
    i=0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("## "):
            current_section = line[3:].strip()
            sections.append(current_section)
            i += 1; continue
        m = ENTRY_RE.match(line)
        if not m:
            i += 1; continue
        vid=m.group(1)
        if vid in out: raise ValueError(f"Duplicate Voice ID in script: {vid}")
        e=ScriptEntry(vid,m.group(2).strip(),section=current_section)
        i += 1
        while i < len(lines):
            meta=lines[i].rstrip()
            if meta.startswith("Type:"):
                e.voice_type=meta.split(":",1)[1].strip(); i+=1; continue
            if meta.startswith("Speaker:"):
                e.speaker=meta.split(":",1)[1].strip(); i+=1; continue
            if meta.startswith("Estimated Duration:"):
                e.duration=meta.split(":",1)[1].strip(); i+=1; continue
            if meta.strip()=="```performance":
                i+=1; body=[]
                while i<len(lines) and lines[i].strip()!="```":
                    body.append(lines[i].rstrip()); i+=1
                if i>=len(lines): raise ValueError(f"Unclosed performance block for {vid}")
                e.performance="\n".join(body).strip(); i+=1; break
            if meta.startswith("### ") or meta.startswith("## "):
                break
            i+=1
        if not current_section: raise ValueError(f"{vid} appears before a section")
        for k,v in {
            "Type":e.voice_type,
            "Speaker":e.speaker,
            "Estimated Duration":e.duration,
            "Performance Script":e.performance,
        }.items():
            if not v: raise ValueError(f"{vid} missing {k}")
        if not has_initial_performance_tag(e.performance):
            raise ValueError(
                f"{vid} performance must begin with at least one initial [performance direction] tag"
            )
        out[vid]=e
    if not out: raise ValueError("No Voice IDs found in script")
    return sections,out


def validate_state(path: Path) -> str:
    text=path.read_text(encoding="utf-8")
    m=STATUS_RE.search(text)
    if not m: raise ValueError("voice-state.yaml missing status")
    status=m.group(1)
    allowed = {"voice_script_ready", "voice_validation", "needs_revision", "voice_delivery_ready"}
    if status not in allowed:
        raise ValueError(f"Flow 7 cannot validate status {status}; expected one of: {', '.join(sorted(allowed))}")
    return status


def validate_project_html(
    path: Path,
    sections: list[str],
    script: dict[str, ScriptEntry],
    requirements: dict[str, Requirement],
) -> list[str]:
    source = path.read_text(encoding="utf-8")
    issues=[]
    if 'id="production-assets-style"' not in source:
        issues.append("Project HTML missing Production Assets presentation")
    if "Production Assets" not in source or "production-assets-nav" not in source:
        issues.append("Project HTML missing Production Assets Voice navigation")
    if source.count('class="pa-row pa-row-voice"') != len(script):
        issues.append("Project HTML compact Voice row count differs from canonical script")

    for section in sections:
        plain = re.sub(r"^\s*\d+\.\s*", "", section).strip()
        if plain and html.escape(plain, quote=True) not in source:
            issues.append(f"Project HTML missing Voice gameplay section: {plain}")

    for vid,e in script.items():
        prompt_id=f"voice-prompt-{vid.lower()}"
        pattern=re.compile(rf'<pre class="voice-script-text" id="{re.escape(prompt_id)}">(.*?)</pre>', re.S)
        matches=pattern.findall(source)
        if len(matches)!=1:
            issues.append(f"Project HTML must contain exact Voice prompt panel once for {vid}")
            continue
        actual=html.unescape(matches[0])
        if actual != e.performance:
            issues.append(f"Project HTML performance text differs from canonical script for {vid}")
        identity = html.escape(f"{e.speaker} — {e.title}", quote=True)
        if identity not in source:
            issues.append(f"Project HTML missing compact Voice identity for {vid}")
    return issues


def validate_docx(path: Path, sections: list[str], script: dict[str, ScriptEntry]) -> list[str]:
    doc=Document(path)
    paras=[p.text for p in doc.paragraphs]
    full="\n".join(paras)
    flat=norm(full)
    issues=[]
    for section in sections:
        if section and norm(section) not in flat:
            issues.append(f"DOCX missing section heading: {section}")
    expected=set(script)
    actual=set(VOICE_ID_RE.findall(full))
    if actual != expected:
        missing=sorted(expected-actual); extra=sorted(actual-expected)
        if missing: issues.append("DOCX missing Voice IDs: "+", ".join(missing))
        if extra: issues.append("DOCX has extra Voice IDs: "+", ".join(extra))
    for vid,e in script.items():
        if full.count(vid) != 1:
            issues.append(f"DOCX must contain {vid} exactly once")
        if f"Speaker: {e.speaker}" not in full:
            issues.append(f"DOCX missing speaker for {vid}: {e.speaker}")
        if e.duration not in full:
            issues.append(f"DOCX missing duration for {vid}: {e.duration}")
        if norm(e.performance) not in flat:
            issues.append(f"DOCX performance text differs/missing for {vid}")
    sec=doc.sections[0]
    w=round(sec.page_width.inches,2); h=round(sec.page_height.inches,2)
    if (w,h)!=(8.5,11.0): issues.append(f"DOCX page size expected Letter 8.5x11, got {w}x{h}")
    if not doc.paragraphs: issues.append("DOCX contains no paragraphs")
    return issues


def main() -> int:
    ap=argparse.ArgumentParser(description="Mechanically validate current Voice requirements/script and any derived project HTML/DOCX presentation.")
    ap.add_argument("project",type=Path, help="workspace/active/<project> directory")
    args=ap.parse_args(); p=args.project
    req=p/"work/voice-requirements.md"; scr=p/"work/voice-production.md"; state=p/"state/voice-state.yaml"
    state_text=state.read_text(encoding="utf-8")
    html_match=re.search(r"(?m)^\s*project_html:\s*(.*?)\s*$", state_text)
    html_ref=html_match.group(1).strip() if html_match else ""
    html_path=(p/html_ref) if html_ref else None
    docx=p/"output/Voice Production.docx"
    missing=[str(x.relative_to(p)) for x in (req,scr,state) if not x.is_file()]
    if missing:
        print("VOICE VALIDATION FAILED: missing files: "+", ".join(missing),file=sys.stderr); return 2
    try:
        validate_state(state)
        requirements=parse_requirements(req)
        sections,script=parse_script(scr)
        issues=validate_revision_identity(p,req,scr,state)
        if set(requirements)!=set(script):
            mi=sorted(set(requirements)-set(script)); ex=sorted(set(script)-set(requirements))
            if mi: issues.append("Script missing Voice IDs: "+", ".join(mi))
            if ex: issues.append("Script has extra Voice IDs: "+", ".join(ex))
        for vid in sorted(set(requirements)&set(script)):
            if requirements[vid].voice_type.casefold()!=script[vid].voice_type.casefold():
                issues.append(f"Type mismatch for {vid}")
            if requirements[vid].speaker.casefold()!=script[vid].speaker.casefold():
                issues.append(f"Speaker mismatch for {vid}")
        if html_path is not None and html_path.is_file():
            issues.extend(validate_project_html(html_path,sections,script,requirements))
        if docx.is_file():
            issues.extend(validate_docx(docx,sections,script))
        if issues:
            print("VOICE VALIDATION FAILED")
            for x in issues: print("- "+x)
            return 1
        print("VOICE VALIDATION PASS")
        print(f"requirements={len(requirements)} script_entries={len(script)} sections={len(sections)}")
        print("project_html="+("passed" if html_path is not None and html_path.is_file() else "not_provided"))
        print("docx="+("passed" if docx.is_file() else "optional_not_provided"))
        print("semantic_and_visual_review=required")
        return 0
    except (OSError,ValueError) as e:
        print(f"VOICE VALIDATION FAILED: {e}",file=sys.stderr); return 2
if __name__=="__main__": raise SystemExit(main())
