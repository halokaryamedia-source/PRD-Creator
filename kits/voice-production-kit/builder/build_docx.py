#!/usr/bin/env python3
"""Build reference-styled Voice Production.docx from canonical Flow 6 Markdown."""
from __future__ import annotations

import argparse
import hashlib
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = "2F65A7"
DARK = "1F2937"
GRAY = "6B7280"
MAIN_FILL = "EAF1FA"
OTHER_FILL = "F4F6F8"
PLACEHOLDER_RE = re.compile(r"\b(?:TBD|TODO|FIXME)\b|\[OPEN\]", re.I)
ENTRY_RE = re.compile(r"^###\s+([A-Za-z0-9][A-Za-z0-9-]*)\s+[—-]\s+(.+?)\s*$")
REQ_ENTRY_RE = re.compile(r"^###\s+([A-Za-z0-9][A-Za-z0-9-]*)\s+[—-]\s+(.+?)\s*$")
SHA256_RE = re.compile(r"^[0-9a-fA-F]{64}$")


@dataclass
class VoiceEntry:
    voice_id: str
    title: str
    voice_type: str = ""
    duration: str = ""
    performance: str = ""


@dataclass
class VoiceSection:
    title: str
    entries: list[VoiceEntry] = field(default_factory=list)


@dataclass
class VoiceDocument:
    title: str
    version: str
    source_requirements: str
    source_requirements_sha256: str
    sections: list[VoiceSection]


def text_fingerprint(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def docx_revision_identifier(requirements_sha256: str, script_sha256: str) -> str:
    return (
        f"voice-requirements-sha256={requirements_sha256};"
        f"voice-script-sha256={script_sha256}"
    )


def parse_script(path: Path) -> VoiceDocument:
    text = path.read_text(encoding="utf-8")
    if PLACEHOLDER_RE.search(text):
        raise ValueError("Canonical voice script still contains an unresolved placeholder.")

    lines = text.splitlines()
    title = ""
    version = "1.0"
    source_requirements = ""
    source_requirements_sha256 = ""
    sections: list[VoiceSection] = []
    current_section: VoiceSection | None = None
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("# ") and not title:
            title = line[2:].strip()
            i += 1
            continue
        if line.startswith("Version:"):
            version = line.split(":", 1)[1].strip() or version
            i += 1
            continue
        if line.startswith("Source Voice Requirements:"):
            if source_requirements:
                raise ValueError("Canonical voice script contains duplicate Source Voice Requirements metadata.")
            source_requirements = line.split(":", 1)[1].strip()
            i += 1
            continue
        if line.startswith("Source Voice Requirements SHA-256:"):
            if source_requirements_sha256:
                raise ValueError("Canonical voice script contains duplicate Voice Requirements SHA-256 metadata.")
            source_requirements_sha256 = line.split(":", 1)[1].strip().lower()
            if not SHA256_RE.fullmatch(source_requirements_sha256):
                raise ValueError("Source Voice Requirements SHA-256 must be exactly 64 hexadecimal characters.")
            i += 1
            continue
        if line.startswith("## "):
            current_section = VoiceSection(line[3:].strip())
            sections.append(current_section)
            i += 1
            continue

        m = ENTRY_RE.match(line)
        if m:
            if current_section is None:
                raise ValueError(f"Voice entry {m.group(1)} appears before a gameplay section.")
            entry = VoiceEntry(m.group(1), m.group(2).strip())
            i += 1
            while i < len(lines):
                meta = lines[i].rstrip()
                if meta.startswith("Type:"):
                    entry.voice_type = meta.split(":", 1)[1].strip()
                    i += 1
                    continue
                if meta.startswith("Estimated Duration:"):
                    entry.duration = meta.split(":", 1)[1].strip()
                    i += 1
                    continue
                if meta.strip() == "```performance":
                    i += 1
                    body: list[str] = []
                    while i < len(lines) and lines[i].strip() != "```":
                        body.append(lines[i].rstrip())
                        i += 1
                    if i >= len(lines):
                        raise ValueError(f"Unclosed performance block for {entry.voice_id}.")
                    entry.performance = "\n".join(body).strip()
                    i += 1
                    break
                if meta.startswith("### ") or meta.startswith("## "):
                    break
                i += 1
            if not entry.voice_type:
                raise ValueError(f"{entry.voice_id} is missing Type.")
            if not entry.duration:
                raise ValueError(f"{entry.voice_id} is missing Estimated Duration.")
            if not entry.performance:
                raise ValueError(f"{entry.voice_id} has an empty Performance Script.")
            current_section.entries.append(entry)
            continue
        i += 1

    if not title:
        raise ValueError("Document requires one '# <Project> Voice Production' heading.")
    if not source_requirements:
        raise ValueError("Canonical voice script requires Source Voice Requirements metadata.")
    if not source_requirements_sha256:
        raise ValueError("Canonical voice script requires Source Voice Requirements SHA-256 metadata.")
    if not sections or not any(section.entries for section in sections):
        raise ValueError("No voice entries were found.")
    ids = [entry.voice_id for section in sections for entry in section.entries]
    if len(ids) != len(set(ids)):
        raise ValueError("Duplicate Voice IDs exist in canonical script.")
    return VoiceDocument(
        title,
        version,
        source_requirements,
        source_requirements_sha256,
        sections,
    )


def parse_requirements(path: Path) -> dict[str, str]:
    result: dict[str, str] = {}
    current_id: str | None = None
    for line in path.read_text(encoding="utf-8").splitlines():
        m = REQ_ENTRY_RE.match(line.rstrip())
        if m:
            current_id = m.group(1)
            if current_id in result:
                raise ValueError(f"Duplicate Voice ID in requirements: {current_id}")
            result[current_id] = ""
            continue
        if current_id and line.startswith("- Type:"):
            result[current_id] = line.split(":", 1)[1].strip()
    if not result:
        raise ValueError("No Voice IDs found in Flow 5 requirements.")
    missing_type = [k for k, v in result.items() if not v]
    if missing_type:
        raise ValueError(f"Requirements missing Type for: {', '.join(missing_type)}")
    return result


def validate_parity(doc: VoiceDocument, requirements: dict[str, str]) -> None:
    actual = {entry.voice_id: entry.voice_type for section in doc.sections for entry in section.entries}
    req_ids, actual_ids = set(requirements), set(actual)
    if req_ids != actual_ids:
        missing = sorted(req_ids - actual_ids)
        extra = sorted(actual_ids - req_ids)
        bits = []
        if missing:
            bits.append("missing script IDs: " + ", ".join(missing))
        if extra:
            bits.append("extra script IDs: " + ", ".join(extra))
        raise ValueError("Voice requirement parity failed: " + "; ".join(bits))
    mismatches = [vid for vid in sorted(req_ids) if requirements[vid].casefold() != actual[vid].casefold()]
    if mismatches:
        raise ValueError("Voice Type differs from Flow 5 requirement for: " + ", ".join(mismatches))


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def shade(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_keep(paragraph, keep_next: bool = False, keep_lines: bool = False) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def bottom_border(paragraph, color: str = "4F81BD") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = rgb(BLUE)
    title.paragraph_format.space_after = Pt(15)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Aptos Display"
    heading.font.size = Pt(17)
    heading.font.bold = True
    heading.font.color.rgb = rgb(BLUE)
    heading.paragraph_format.space_before = Pt(24)
    heading.paragraph_format.space_after = Pt(0)


def section_subtitle(section: VoiceSection) -> str:
    types: list[str] = []
    for entry in section.entries:
        if entry.voice_type not in types:
            types.append(entry.voice_type)
    if len(types) == 1:
        return types[0]
    if len(types) == 2:
        return f"{types[0]} and {types[1]}"
    return ", ".join(types[:-1]) + f", and {types[-1]}"


def add_script_paragraph(doc: Document, entry: VoiceEntry) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    set_keep(p, keep_lines=True)
    fill = MAIN_FILL if entry.voice_type.casefold() == "main story" else OTHER_FILL
    shade(p, fill)

    lines = entry.performance.split("\n")
    for n, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if re.fullmatch(r"\s*\[[^\]]+\]\s*", line):
            run.font.italic = True
            run.font.color.rgb = rgb(BLUE)
        else:
            run.font.color.rgb = rgb(DARK)
        if n != len(lines) - 1:
            run.add_break()


def build_docx(
    data: VoiceDocument,
    output: Path,
    requirements_sha256: str,
    script_sha256: str,
) -> None:
    doc = Document()
    configure_styles(doc)
    doc.core_properties.identifier = docx_revision_identifier(
        requirements_sha256,
        script_sha256,
    )

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.70)
    sec.right_margin = Inches(0.70)

    cover = doc.add_paragraph(style="Title")
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run(data.title)
    bottom_border(cover)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Version {data.version}")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ElevenLabs-ready performance scripts")
    r.italic = True
    r.font.color.rgb = rgb(GRAY)

    doc.add_page_break()

    for sidx, section in enumerate(data.sections, 1):
        h = doc.add_paragraph(section.title, style="Heading 1")
        if sidx > 1:
            h.paragraph_format.page_break_before = True
        set_keep(h, keep_next=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_keep(p, keep_next=True)
        r = p.add_run(section_subtitle(section))
        r.italic = True
        r.font.color.rgb = rgb(GRAY)

        for entry in section.entries:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            set_keep(p, keep_next=True)
            r = p.add_run(entry.voice_type.upper())
            r.font.name = "Aptos"
            r.font.size = Pt(8.5)
            r.font.bold = True
            r.font.color.rgb = rgb(BLUE)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            set_keep(p, keep_next=True)
            r = p.add_run(f"{entry.voice_id} - {entry.title}")
            r.font.name = "Aptos Display"
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = rgb(DARK)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            set_keep(p, keep_next=True)
            r = p.add_run(f"Estimated Duration: {entry.duration}")
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = rgb(GRAY)

            add_script_paragraph(doc, entry)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("script", type=Path, help="Canonical work/voice-production.md")
    parser.add_argument("output", type=Path, help="Destination Voice Production.docx")
    parser.add_argument(
        "--requirements",
        type=Path,
        required=True,
        help="Current Flow 5 work/voice-requirements.md",
    )
    args = parser.parse_args()

    try:
        if not args.script.is_file():
            raise FileNotFoundError(args.script)
        if not args.requirements.is_file():
            raise FileNotFoundError(args.requirements)

        data = parse_script(args.script)
        requirements_sha256 = text_fingerprint(args.requirements)
        if data.source_requirements_sha256.casefold() != requirements_sha256.casefold():
            raise ValueError(
                "Canonical voice script Voice Requirements SHA-256 does not match the current requirements file. "
                f"expected {requirements_sha256}, found {data.source_requirements_sha256}"
            )

        validate_parity(data, parse_requirements(args.requirements))
        script_sha256 = text_fingerprint(args.script)
        build_docx(data, args.output, requirements_sha256, script_sha256)
        print(args.output)
        return 0
    except (OSError, ValueError) as exc:
        print(f"VOICE DOCX BUILD FAILED: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
